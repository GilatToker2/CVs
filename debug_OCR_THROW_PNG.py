from anthropic import AnthropicVertex
import base64
import os
import sys
import subprocess
from pathlib import Path
import tempfile
import platform
import shutil
import urllib.request
import zipfile
import time

LOCATION = "us-east5"
PROJECT_ID = "kpmg-project-2"
MODEL = "claude-3-7-sonnet@20250219"


def install_required_packages():
    """Install required packages if they're not already installed."""
    required_packages = [
        'pdf2image',
        'python-docx',
        'Pillow',
        'pywin32',  # For doc to docx conversion
        'tqdm',  # For progress bars
        'PyMuPDF',  # For PDF rendering as fallback
        'reportlab'  # For creating PDFs from scratch
    ]

    for package in required_packages:
        try:
            __import__(package.replace('-', '_').split('[')[0])
            print(f"{package} is already installed.")
        except ImportError:
            print(f"Installing {package}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])
            print(f"{package} has been installed.")

    # Download and install poppler for Windows
    if platform.system() == "Windows":
        download_and_install_poppler()


def download_and_install_poppler():
    """Download and install poppler for Windows."""
    poppler_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "poppler", "bin")
    if not os.path.exists(poppler_path):
        print("Poppler not found. Downloading and installing poppler for Windows...")
        try:
            # Create poppler directory
            os.makedirs("poppler", exist_ok=True)

            # Download poppler
            poppler_url = "https://github.com/oschwartz10612/poppler-windows/releases/download/v23.08.0-0/Release-23.08.0-0.zip"
            zip_path = os.path.join("poppler", "poppler.zip")
            print("Downloading poppler...")
            urllib.request.urlretrieve(poppler_url, zip_path)

            # Extract poppler
            print("Extracting poppler...")
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall("poppler")

            # Clean up
            os.remove(zip_path)
            print("Poppler installed successfully.")
            return True
        except Exception as e:
            print(f"Error installing poppler: {str(e)}")
            print(
                "Please download and install poppler manually from: https://github.com/oschwartz10612/poppler-windows/releases")
            print("Extract it to a folder named 'poppler' in the same directory as this script.")
            return False
    return True


def get_poppler_path():
    """Get the path to poppler binaries."""
    if platform.system() == "Windows":
        # Check for poppler in the script directory
        script_dir = os.path.dirname(os.path.abspath(__file__))
        poppler_path = os.path.join(script_dir, "poppler", "Library", "bin")

        if os.path.exists(poppler_path):
            return poppler_path

        poppler_path = os.path.join(script_dir, "poppler", "bin")

        if os.path.exists(poppler_path):
            return poppler_path

        # Check for poppler in common installation locations
        common_paths = [
            r"C:\Program Files\poppler\bin",
            r"C:\Program Files (x86)\poppler\bin",
            r"C:\poppler\bin"
        ]

        for path in common_paths:
            if os.path.exists(path):
                return path

        return None
    else:
        # On Linux/Mac, poppler should be in PATH
        return None


def convert_pdf_to_png_with_pymupdf(pdf_path):
    """Convert the first page of a PDF to a PNG image using PyMuPDF."""
    try:
        import fitz  # PyMuPDF

        # Open the PDF
        doc = fitz.open(pdf_path)

        # Check if the document has pages
        if doc.page_count == 0:
            print(f"PDF has no pages: {pdf_path}")
            return None

        # Get the first page
        page = doc[0]

        # Render page to an image (with higher resolution)
        pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))

        # Create a temporary file for the image
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
            temp_path = temp_file.name

        # Save the image
        pix.save(temp_path)

        print(f"Converted PDF to PNG using PyMuPDF: {temp_path}")
        return temp_path
    except Exception as e:
        print(f"PyMuPDF conversion failed: {str(e)}")
        return None


def convert_pdf_to_png(pdf_path):
    """Convert the first page of a PDF to a PNG image."""
    # First try with pdf2image and poppler
    try:
        from pdf2image import convert_from_path

        poppler_path = get_poppler_path()

        # Convert the first page of the PDF to an image
        if poppler_path:
            print(f"Using poppler from: {poppler_path}")
            images = convert_from_path(
                pdf_path,
                first_page=1,
                last_page=1,
                poppler_path=poppler_path,
                dpi=300  # Higher DPI for better quality
            )
        else:
            # Try without specifying poppler path (for Linux/Mac)
            images = convert_from_path(
                pdf_path,
                first_page=1,
                last_page=1,
                dpi=300
            )

        if images:
            # Create a temporary file for the image
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                temp_path = temp_file.name

            # Save the image to the temporary file
            images[0].save(temp_path, 'PNG')

            print(f"Converted PDF to PNG with pdf2image: {temp_path}")
            return temp_path
        else:
            print(f"Failed to convert PDF to PNG with pdf2image: {pdf_path}")
            # Fall through to next method
    except Exception as e:
        print(f"Error converting PDF to PNG with pdf2image: {str(e)}")
        # Fall through to next method

    # Try with PyMuPDF
    png_path = convert_pdf_to_png_with_pymupdf(pdf_path)
    if png_path:
        return png_path

    # Try with ImageMagick if available
    try:
        print("Attempting conversion with ImageMagick...")
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
            temp_path = temp_file.name

        # Try using ImageMagick's convert command
        subprocess.run(
            ["magick", "convert", "-density", "300", f"{pdf_path}[0]", temp_path],
            check=True
        )

        print(f"Converted PDF to PNG using ImageMagick: {temp_path}")
        return temp_path
    except Exception as img_err:
        print(f"ImageMagick conversion failed: {str(img_err)}")

    # If all methods failed, return None
    return None


def create_text_image(text, width=800, height=1000):
    """Create a PNG image with text content."""
    try:
        from PIL import Image, ImageDraw, ImageFont

        # Create a white image
        image = Image.new('RGB', (width, height), color='white')
        draw = ImageDraw.Draw(image)

        # Try to use a default font
        try:
            if platform.system() == "Windows":
                font_path = os.path.join(os.environ['WINDIR'], 'Fonts', 'Arial.ttf')
                font = ImageFont.truetype(font_path, 12)
            else:
                font = ImageFont.load_default()
        except Exception:
            font = ImageFont.load_default()

        # Draw the text
        draw.text((10, 10), text, fill='black', font=font)

        # Save to a temporary file
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
            temp_path = temp_file.name

        image.save(temp_path)
        return temp_path
    except Exception as e:
        print(f"Error creating text image: {str(e)}")
        return None


def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file."""
    try:
        import docx

        doc = docx.Document(docx_path)
        full_text = []

        for para in doc.paragraphs:
            full_text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from DOCX: {str(e)}")
        return "Error extracting text from document."


def create_pdf_from_text(text):
    """Create a PDF from text content."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph

        # Create a temporary file for the PDF
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            temp_pdf_path = temp_pdf.name

        # Create the PDF
        doc = SimpleDocTemplate(temp_pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()

        # Split text into paragraphs
        paragraphs = text.split('\n')
        story = [Paragraph(p, styles['Normal']) for p in paragraphs if p.strip()]

        # Build the PDF
        doc.build(story)

        return temp_pdf_path
    except Exception as e:
        print(f"Error creating PDF from text: {str(e)}")
        return None


def convert_docx_to_png(docx_path):
    """Convert a DOCX file to a PNG image."""
    # Method 1: Try using Word COM interface directly
    try:
        import win32com.client

        # Create a temporary file for the PDF
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            temp_pdf_path = temp_pdf.name

        # Initialize Word application
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        try:
            # Open the docx file
            doc = word.Documents.Open(os.path.abspath(docx_path))

            # Save as PDF
            doc.SaveAs(os.path.abspath(temp_pdf_path), FileFormat=17)  # 17 represents PDF format
            doc.Close()

            # Convert PDF to PNG
            png_path = convert_pdf_to_png(temp_pdf_path)

            # Clean up the temporary PDF file
            try:
                os.remove(temp_pdf_path)
            except:
                pass

            if png_path:
                return png_path
        finally:
            # Make sure Word is closed
            try:
                word.Quit()
            except:
                pass
    except Exception as e:
        print(f"Word COM interface conversion failed: {str(e)}")

    # Method 2: Extract text and create an image
    try:
        print("Extracting text from DOCX and creating image...")
        text = extract_text_from_docx(docx_path)

        if text:
            # Create a PDF from the text
            pdf_path = create_pdf_from_text(text)

            if pdf_path:
                # Convert the PDF to PNG
                png_path = convert_pdf_to_png(pdf_path)

                # Clean up
                try:
                    os.remove(pdf_path)
                except:
                    pass

                if png_path:
                    return png_path

            # If PDF conversion failed, create a direct image from text
            return create_text_image(text)
    except Exception as e:
        print(f"Text extraction and image creation failed: {str(e)}")

    # If all methods failed, return None
    return None


def extract_text_from_doc(doc_path):
    """Extract text from a DOC file using a fallback method."""
    try:
        # Try to use antiword if available
        try:
            result = subprocess.run(
                ["antiword", doc_path],
                capture_output=True,
                text=True,
                check=True
            )
            return result.stdout
        except:
            pass

        # Try to use catdoc if available
        try:
            result = subprocess.run(
                ["catdoc", doc_path],
                capture_output=True,
                text=True,
                check=True
            )
            return result.stdout
        except:
            pass

        # If all else fails, return a message
        return "Could not extract text from DOC file. Please convert to DOCX or PDF."
    except Exception as e:
        print(f"Error extracting text from DOC: {str(e)}")
        return "Error extracting text from document."


def convert_doc_to_png(doc_path):
    """Convert a DOC file to a PNG image."""
    # Method 1: Try using Word COM interface directly
    try:
        import win32com.client

        # Create a temporary file for the PDF
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            temp_pdf_path = temp_pdf.name

        # Initialize Word application
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        try:
            # Open the doc file
            doc = word.Documents.Open(os.path.abspath(doc_path))

            # Save as PDF
            doc.SaveAs(os.path.abspath(temp_pdf_path), FileFormat=17)  # 17 represents PDF format
            doc.Close()

            # Convert PDF to PNG
            png_path = convert_pdf_to_png(temp_pdf_path)

            # Clean up the temporary PDF file
            try:
                os.remove(temp_pdf_path)
            except:
                pass

            if png_path:
                return png_path
        finally:
            # Make sure Word is closed
            try:
                word.Quit()
            except:
                pass
    except Exception as e:
        print(f"Word COM interface conversion failed: {str(e)}")

    # Method 2: Try converting to DOCX first
    try:
        import win32com.client

        # Create a temporary file for the DOCX
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            temp_docx_path = temp_docx.name

        # Initialize Word application
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        try:
            # Open the doc file
            doc = word.Documents.Open(os.path.abspath(doc_path))

            # Save as docx
            doc.SaveAs(os.path.abspath(temp_docx_path), 16)  # 16 represents the wdFormatDocumentDefault (docx)
            doc.Close()

            # Convert DOCX to PNG
            png_path = convert_docx_to_png(temp_docx_path)

            # Clean up the temporary DOCX file
            try:
                os.remove(temp_docx_path)
            except:
                pass

            if png_path:
                return png_path
        finally:
            # Make sure Word is closed
            try:
                word.Quit()
            except:
                pass
    except Exception as e:
        print(f"DOC to DOCX conversion failed: {str(e)}")

    # Method 3: Extract text and create an image
    try:
        print("Extracting text from DOC and creating image...")
        text = extract_text_from_doc(doc_path)

        if text:
            # Create a PDF from the text
            pdf_path = create_pdf_from_text(text)

            if pdf_path:
                # Convert the PDF to PNG
                png_path = convert_pdf_to_png(pdf_path)

                # Clean up
                try:
                    os.remove(pdf_path)
                except:
                    pass

                if png_path:
                    return png_path

            # If PDF conversion failed, create a direct image from text
            return create_text_image(text)
    except Exception as e:
        print(f"Text extraction and image creation failed: {str(e)}")

    # If all methods failed, return None
    return None


def convert_file_to_png(file_path):
    """Convert any supported file to PNG."""
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()

    if file_extension == '.pdf':
        return convert_pdf_to_png(file_path)
    elif file_extension == '.docx':
        return convert_docx_to_png(file_path)
    elif file_extension == '.doc':
        return convert_doc_to_png(file_path)
    elif file_extension in ['.png', '.jpg', '.jpeg', '.gif', '.webp']:
        # Already an image, no need to convert
        return file_path
    else:
        print(f"Unsupported file type for conversion: {file_extension}")
        return None


def extract_text_with_claude(file_path, prompt_text="Please extract all text from this document."):
    """
    Send a PNG image to Claude and get the extracted text.

    Args:
        file_path: Path to the PNG file to process
        prompt_text: Text to send to Claude along with the file

    Returns:
        The text response from Claude
    """
    try:
        # Initialize the client
        client = AnthropicVertex(region=LOCATION, project_id=PROJECT_ID)

        # Read the file as binary
        with open(file_path, 'rb') as file:
            file_content = file.read()

        # Encode the file content as base64
        base64_content = base64.b64encode(file_content).decode('utf-8')

        # Create the message with the file content
        message = client.messages.create(
            max_tokens=4096,  # Increased token limit for longer responses
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": prompt_text
                        },
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": "image/png",
                                "data": base64_content
                            }
                        }
                    ]
                }
            ],
            model=MODEL
        )

        # Extract and return the text response
        return message.content[0].text

    except Exception as e:
        return f"Error processing file with Claude: {str(e)}"


def process_client_files_folder(folder_path):
    """Process all files in the client files folder using Claude."""
    folder = Path(folder_path)

    if not folder.exists():
        print(f"Folder '{folder_path}' does not exist.")
        return

    files = list(folder.iterdir())
    files = [f for f in files if f.is_file() and not f.name.startswith('~$')]  # Skip temporary files

    if not files:
        print(f"No files found in '{folder_path}'.")
        return

    print(f"Found {len(files)} files in '{folder_path}'.")

    # Import tqdm for progress bar
    try:
        from tqdm import tqdm
        files_iter = tqdm(files)
    except ImportError:
        files_iter = files

    for file_path in files_iter:
        print(f"\n{'=' * 50}")
        print(f"Processing: {file_path.name}")
        print(f"{'=' * 50}")

        # Convert file to PNG
        png_path = convert_file_to_png(file_path)

        if png_path:
            # Extract text using Claude
            text = extract_text_with_claude(png_path)

            # Print the extracted text
            print(text)

            # Clean up temporary PNG file if it's different from the original file
            if png_path != str(file_path) and os.path.exists(png_path):
                try:
                    os.remove(png_path)
                    print(f"Removed temporary file: {png_path}")
                except:
                    pass
        else:
            print(f"Failed to convert {file_path} to PNG")

        print(f"{'=' * 50}\n")


if __name__ == "__main__":
    # Install required packages
    # install_required_packages()

    # Process the client files folder
    client_files_folder = "קבצי קוח"
    process_client_files_folder(client_files_folder)
